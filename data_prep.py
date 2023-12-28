# Install necessary packages

%pip install python-docx pdfplumber spacy langchain tiktoken
!python -m spacy download en_core_web_sm

dbutils.library.restartPython()

import re
import json
import time
import requests
import spacy
from docx import Document
import pdfplumber
from langchain.embeddings import AzureOpenAIEmbeddings


# Extracts regular docx files and chunks them by their headers into a list

def is_header_docx(paragraph):
    # Check if the paragraph style is a header style
    if paragraph.style:
        if paragraph.style.name.startswith('Heading'):
            return True
        
    # Check if the paragraph is italicized/underlined/bold and ends with ':'
    if paragraph.text.endswith(':'):
        if any(run.bold for run in paragraph.runs) or \
           any(run.italic for run in paragraph.runs) or \
           any(run.underline for run in paragraph.runs):
            return True

    return False

def extract_docx(path):
    # Load the document
    doc = Document(path)
    content = []
    current_header = ""
    header_content = ""

    # Iterate through all paragraphs in the document
    for para in doc.paragraphs:
        if is_header_docx(para):
            # Process the accumulated text if any exists
            if header_content:
                # Split the text into sentences and append to content list
                sentences = re.split(r'(?<=[.!?])\s+', header_content)
                for sentence in sentences:
                    content.append(f'{current_header} - {sentence}')

            # Update the current header and reset the header content
            current_header = para.text.strip()
            header_content = ""
        else:
            # Accumulate text under the current header
            header_content += para.text + " "

    # Process any remaining text after the last paragraph
    if header_content:
        sentences = re.split(r'(?<=[.!?])\s+', header_content)
        for sentence in sentences:
            content.append(f'{current_header} - {sentence}')

    return content

bug_reporting_text = extract_docx(MOUNT_LOCATION + "/Bug Reporting Directions.docx")


# Specialized extraction for Docx Files with tables (was not extracting properly with above function)

def extract_docx_tables(path):
    # Load the document
    doc = Document(path)
    content = []

    # Iterate through all elements in the document body
    for element in doc.element.body:
        # Check if the current element is a paragraph
        if element.tag.endswith('p'):
            # Add the paragraph text to the content list, stripping leading/trailing whitespace
            content.append(element.text.strip())

        # Check if the current element is a table
        elif element.tag.endswith('tbl'):
            # Iterate through all rows in the table
            for row in element.findall('.//w:tr', doc.element.nsmap):
                row_data = []
                # Iterate through all cells in the current row
                for cell in row.findall('.//w:tc', doc.element.nsmap):
                    # Extract and concatenate text from each XML node inside the cell
                    cell_text = ''.join(node.text for node in cell.findall('.//w:t', doc.element.nsmap))
                    row_data.append(cell_text.strip())

                content.append(' | '.join(row_data))

    return content

federal_holidays_text = extract_docx_tables(MOUNT_LOCATION + "/Federal Holidays 2023 - 2024.docx")


# Extracts pdf files and chunks them by their headers into a list

def is_header_pdf(line):
    # Enhanced header check: includes numbered headers and specific patterns
    numbered_header_pattern = r'^(?:[IVXLCDM]+\.\s+|[A-Z]+\.)'
    if re.match(numbered_header_pattern, line) or line.isupper() or line.istitle():
        return line
    return None

def is_end_of_sentence(line):
    return re.search(r'[.!?]$', line.strip()) is not None

def extract_pdf(path):
    sections = []
    current_header = ""
    header_content = ""

    with pdfplumber.open(path) as pdf:
        # Iterate through each page in the PDF
        for page in pdf.pages:
            page_text = page.extract_text()
            # Skip the page if no text is found
            if not page_text:
                continue

            # Split the page text into lines
            lines = page_text.split('\n')
            for line in lines:
                line = line.strip()
                # Check if the line is a header
                if is_header_pdf(line):
                    if current_header and header_content:
                        # Split the text into sentences and add to the sections list
                        sentences = re.split(r'(?<=[.!?])\s+', header_content)
                        for sentence in sentences:
                            sections.append(f'{current_header} - {sentence}')

                    # Update the current header and reset the text buffer
                    current_header = line
                    header_content = ""
                elif current_header is not None:
                    if not is_end_of_sentence(header_content) and header_content:
                        header_content += ' ' + line
                    else:
                        header_content += line

            # Process any remaining text after the last line on the page
            if current_header and header_content:
                sentences = re.split(r'(?<=[.!?])\s+', header_content)
                for sentence in sentences:
                    sections.append(f'{current_header} - {sentence}')

    return sections

benefit_guide_text = extract_pdf(MOUNT_LOCATION + "/Sample Benefit Guide.pdf")
employee_handbook_text = extract_pdf(MOUNT_LOCATION + "/Sample Employee Handbook.pdf")
training_policy_text = extract_pdf(MOUNT_LOCATION + "/Sample Training Development Policy.pdf")


# Cleans the files and then resaves as a list

nlp = spacy.load("en_core_web_sm")

def clean_text(text):
    text = text.replace("’", "").replace("'", "")
    doc = nlp(text)
    cleaned_text = []

    for token in doc:
        # Lemmatize and lower the token
        lemmatized_token = token.lemma_.lower()

        # Check if the token is not space, not empty, not a vertical bar, not an apostrophe,
        # and is either a period, a hyphen, or non-punctuation
        if (not token.is_space and lemmatized_token and lemmatized_token not in ["|", "’", "'"] and 
            (lemmatized_token == '.' or lemmatized_token == '-' or not token.is_punct)):
            cleaned_text.append(lemmatized_token)

    # Join the cleaned tokens into a single string
    result = " ".join(cleaned_text)
    # Correct spaces before periods if any
    result = re.sub(r'\s+\.', '.', result)
    return result

def remove_duplicates(input_list):
    unique_list = []
    for item in input_list:
        if item not in unique_list:
            unique_list.append(item)
    return unique_list

bug_reporting_cleaned = remove_duplicates([clean_text(text) for text in bug_reporting_text if text.strip() != ''])
federal_holidays_cleaned = remove_duplicates([clean_text(text) for text in federal_holidays_text if text.strip() != ''])
benefit_guide_cleaned = remove_duplicates([clean_text(text) for text in benefit_guide_text if text.strip() != ''])
employee_handbook_cleaned = remove_duplicates([clean_text(text) for text in employee_handbook_text if text.strip() != ''])
training_policy_cleaned = remove_duplicates([clean_text(text) for text in training_policy_text if text.strip() != ''])


# Initialize OpenAIEmbeddings

embeddings_model = AzureOpenAIEmbeddings(
    azure_deployment="embedding",
    model="text-embedding-ada-002",
    azure_endpoint=OPENAI_ENDPOINT,
    openai_api_type="azure",
    openai_api_key=OPENAI_API_KEY,
    openai_api_version="2023-03-15-preview"
)


# Function to embed a list of text chunks

def embed_chunks(chunks, embeddings_model, delay=0):
    embeddings = []
    for chunk in chunks:
        if chunk:
            embeddings.append(embeddings_model.embed_query(chunk))
            time.sleep(delay)
    return embeddings

# Embedding your text chunks
bug_reporting_embeddings = embed_chunks(bug_reporting_cleaned, embeddings_model)
federal_holidays_embeddings = embed_chunks(federal_holidays_cleaned, embeddings_model)
benefit_guide_embeddings = embed_chunks(benefit_guide_cleaned, embeddings_model)
employee_handbook_embeddings = embed_chunks(employee_handbook_cleaned, embeddings_model)
training_policy_embeddings = embed_chunks(training_policy_cleaned, embeddings_model)


# Combine all the embeddings and text chunks

all_embeddings = bug_reporting_embeddings + federal_holidays_embeddings + benefit_guide_embeddings + employee_handbook_embeddings + training_policy_embeddings
all_text_chunks = bug_reporting_cleaned + federal_holidays_cleaned + benefit_guide_cleaned + employee_handbook_cleaned + training_policy_cleaned


# Creates an embeddings list with the content from the text chunks

documents = []
for idx, (embedding, text_chunk) in enumerate(zip(all_embeddings, all_text_chunks)):
    document = {
        "id": idx,  # A unique identifier
        "embedding": json.dumps(embedding),  # The numerical embedding (might need flattening or serialization)
        "content": text_chunk  # The corresponding readable text chunk
    }
    documents.append(document)


# Initialize Azure AI Search

service_name = AI_SEARCH_SERVICE_NAME
index_name = AI_SEARCH_INDEX_NAME
api_key = AI_SEARCH_KEY
api_version = "2023-11-01"  # Updated to the latest API version


# Send embeddings to Azure AI Search Data Index

# Endpoint URL for the Azure Search service
endpoint = f"https://{service_name}.search.windows.net/indexes/{index_name}/docs/index?api-version={api_version}"

# Preparing the payload for the POST request
# This payload consists of a list of documents to be indexed
payload = {
    "value": [
        {
            "@search.action": "upload", # Action type, here it's 'upload'
            "id": str(document["id"]),  # Convert 'id' to string
            "embedding": document["embedding"],  # String
            "content": document["content"]  # String
        } for document in documents
    ]
}

# Headers for the HTTP request
headers = {
    "Content-Type": "application/json",
    "api-key": api_key
}

def send_batch(endpoint, headers, batch):
    # Prepare the payload with the batch of documents
    payload = {"value": batch}

    # Send a POST request to the Azure Search service
    response = requests.post(endpoint, headers=headers, json=payload)

    # Check the response status and print appropriate messages
    if response.status_code in (200, 201):
        print("Batch successfully indexed.")
    else:
        print(f"Error in batch: {response.status_code} - {response.text}")

# Splitting the documents into batches and sending each batch
batch_size = 250  # Adjust this based on Azure AI Search limits
for i in range(0, len(documents), batch_size):
    batch = [
        {
            "@search.action": "upload",
            "id": str(doc["id"]),
            "embedding": doc["embedding"],
            "content": doc["content"]
        } for doc in documents[i:i+batch_size]
    ]
    # Sending the batch for indexing
    send_batch(endpoint, headers, batch)
